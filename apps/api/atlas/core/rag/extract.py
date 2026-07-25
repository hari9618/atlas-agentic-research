"""Extract plain text from uploaded documents.

Uploads used to be text-only (.md/.txt), but the documents people actually have —
filings, reports, briefs — are PDFs and Word files. This turns those into the plain
text the chunker expects. Parsing (not raw bytes) is the point: indexing a PDF's
binary would fill the vector store with garbage, so unsupported types are rejected
rather than stored.

Extractors are optional imports: if a parser isn't installed the caller gets a clear
error for that type instead of an import crash at module load.
"""

from __future__ import annotations

import io
import logging

log = logging.getLogger("atlas.rag.extract")

# Extensions we can turn into text. Kept in one place so the API and UI agree.
SUPPORTED_SUFFIXES = (".txt", ".md", ".markdown", ".pdf", ".docx")


def _from_pdf(raw: bytes) -> str:
    import pdfplumber

    out: list[str] = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                out.append(text)
    return "\n\n".join(out)


def _from_docx(raw: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables often hold the financials in these documents — don't drop them.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(filename: str, raw: bytes) -> tuple[str, str | None]:
    """Return (text, error). Exactly one is meaningful.

    On success ``error`` is None; on failure ``text`` is empty and ``error`` explains
    why (unsupported type, unreadable encoding, empty document, missing parser).
    """
    name = (filename or "").lower()

    try:
        if name.endswith((".txt", ".md", ".markdown")):
            try:
                text = raw.decode("utf-8")
            except UnicodeDecodeError:
                return "", "Text file must be UTF-8 encoded."
        elif name.endswith(".pdf"):
            text = _from_pdf(raw)
        elif name.endswith(".docx"):
            text = _from_docx(raw)
        else:
            return "", f"Unsupported file type. Allowed: {', '.join(SUPPORTED_SUFFIXES)}"
    except ImportError as exc:  # pragma: no cover - parser not installed
        log.warning("parser missing for %s: %s", name, exc)
        return "", "This file type needs an extra parser that isn't installed."
    except Exception as exc:
        log.warning("failed to extract %s: %s", name, exc)
        return "", "Could not read this document — it may be scanned, encrypted, or corrupt."

    if len(text.strip()) < 50:
        return "", "No extractable text found (a scanned/image-only PDF has no text layer)."
    return text, None
